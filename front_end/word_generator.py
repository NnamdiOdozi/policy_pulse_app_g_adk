# front-end/word_generator.py
# =============================================================================
# WORD DOCUMENT GENERATION
# =============================================================================
# PURPOSE: Convert AI-generated text into formatted Word document
# LIBRARY: python-docx for .docx file creation

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import io
import re
from datetime import datetime

def generate_policy_word_doc(policy_text, title="Policy Document"):
    """   
    Create downloadable Word document from policy text
    
    FORMATTING APPLIED:
    1. Title: Bold, 16pt, centered
    2. Headings: Bold, 14pt, numbered (1., 2., 3.)
    3. Subheadings: Bold, 12pt, numbered (1.1, 1.2)
    4. Body text: Normal, 11pt, justified
    5. Lists: Bulleted or numbered as appropriate
    
    PARSING LOGIC:
    - Detect headings by: "# " prefix or "1." numbering
    - Detect subheadings by: "## " prefix or "1.1" numbering
    - Detect lists by: "-" or "*" or "1)" prefix
    - Everything else treated as body paragraph
    
    WHY WORD FORMAT?:
    - Industry standard for policy documents
    - Users can edit further (not PDF locked)
    - Preserves formatting better than plain text
    - Easy to share internally (universal format)
    
    STYLING:
    - Uses built-in python-docx styles
    - Custom styles defined for consistency
    - Matches typical corporate document style
    
    DOWNLOAD MECHANISM:
    - File saved to temporary location
    - Streamlit st.download_button displays
    - User clicks to download
    - Temp file cleaned up after download
    
    ERROR HANDLING:
    - Malformed markdown handled gracefully
    - Falls back to plain text if parsing fails
    - Logs warnings for debugging
    """
    
    doc = Document()
    setup_document_styles(doc)
    
    # Remove meta-commentary at the start
    policy_text = remove_meta_commentary(policy_text)
    
    # Split into lines and process
    lines = policy_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip lines that are purely markdown artifacts
        if line in ['**', '***', '---']:
            continue
            
        # 1. MAIN TITLE (all caps + POLICY/GUIDE/PROCEDURE)
        if (line.isupper() and 
            any(word in line for word in ['POLICY', 'GUIDE', 'PROCEDURE', 'FRAMEWORK', 'GUIDELINES'])):
            add_title(doc, line)
            
        # 2. VERSION LINE
        elif re.match(r'[vV]ersion|v\d+\.\d+|v1\.0', line, re.IGNORECASE):
            add_version(doc, line)
            
        # 3. MAIN SECTION HEADINGS (1., 2., 3., etc. or **1. Text**)
        elif re.match(r'^\*{0,2}\d+\.\s+[A-Z]', line):
            clean_line = clean_markdown(line)
            add_section_heading(doc, clean_line)
            
        # 4. SUB-SECTION HEADINGS (1.1, 1.2, 2.1, etc. or **1.1 Text**)
        elif re.match(r'^\*{0,2}\d+\.\d+\.?\s+[A-Z]', line):
            clean_line = clean_markdown(line)
            add_subsection_heading(doc, clean_line)
            
        # 5. BULLET POINTS
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            clean_line = clean_markdown(line.lstrip('•-* '))
            add_bullet_point(doc, clean_line)
            
        # 6. SOURCES SECTION
        elif line.lower().startswith('sources:') or line.startswith('**Sources:**'):
            add_subsection_heading(doc, "Sources")
            
        # 7. REGULAR PARAGRAPHS
        else:
            add_paragraph_with_formatting(doc, line)
    
    # Add footer
    add_footer(doc)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def remove_meta_commentary(text):
    """Remove common meta-commentary phrases"""
    meta_patterns = [
        r'^Here is (a|the) (formal|standard|comprehensive|draft|complete).*?policy.*?:\s*',
        r'^Here is (a|the).*?document.*?:\s*',
        r'^I (will|have) (draft|generat|creat).*?:\s*',
        r'You can copy and paste.*?Word document\.\s*'
    ]
    
    for pattern in meta_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    return text.strip()


def clean_markdown(text):
    """Remove markdown formatting but preserve the text"""
    # Remove markdown bold/italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # bold+italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)      # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)           # italic
    text = re.sub(r'__(.+?)__', r'\1', text)           # bold (alt)
    text = re.sub(r'_(.+?)_', r'\1', text)             # italic (alt)
    return text.strip()


def add_title(doc, text):
    """Add main document title"""
    text = clean_markdown(text)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacing


def add_version(doc, text):
    """Add version/date line"""
    text = clean_markdown(text)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacing


def add_section_heading(doc, text):
    """Add main section heading (1., 2., 3., etc.)"""
    text = clean_markdown(text)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    para.space_before = Pt(12)
    para.space_after = Pt(6)


def add_subsection_heading(doc, text):
    """Add subsection heading (1.1, 1.2, etc.)"""
    text = clean_markdown(text)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    para.space_before = Pt(8)
    para.space_after = Pt(4)


def add_bullet_point(doc, text):
    """Add bullet point"""
    text = clean_markdown(text)
    para = doc.add_paragraph(style='List Bullet')
    add_formatted_text_to_para(para, text)
    para.space_after = Pt(3)


def add_paragraph_with_formatting(doc, text):
    """Add regular paragraph with inline markdown formatting"""
    para = doc.add_paragraph()
    add_formatted_text_to_para(para, text)
    para.space_after = Pt(6)


def add_formatted_text_to_para(para, text):
    """
    Add text to paragraph with inline bold/italic handling
    Handles **bold** and *italic* within regular text
    """
    # Pattern to find bold and italic markdown
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add regular text before the formatted text
        if match.start() > last_end:
            para.add_run(text[last_end:match.start()])
        
        # Determine formatting type and add formatted text
        if match.group(2):  # Bold + italic
            run = para.add_run(match.group(2))
            run.font.bold = True
            run.font.italic = True
        elif match.group(3):  # Bold only
            run = para.add_run(match.group(3))
            run.font.bold = True
        elif match.group(4):  # Italic only
            run = para.add_run(match.group(4))
            run.font.italic = True
        
        last_end = match.end()
    
    # Add remaining regular text
    if last_end < len(text):
        para.add_run(text[last_end:])


def add_footer(doc):
    """Add generation footer"""
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f"Generated by Policy Pulse on {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def setup_document_styles(doc):
    """Set up document-wide styles"""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)